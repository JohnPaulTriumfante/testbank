from flask import Flask, render_template, request, redirect, send_file, jsonify, url_for
import sqlite3
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
from datetime import datetime
import random
import json
import os
from pathlib import Path

app = Flask(__name__)

# Create directory for storing generated exam PDFs
EXAMS_DIR = Path("generated_exams")
EXAMS_DIR.mkdir(exist_ok=True)

def get_db():
    conn = sqlite3.connect("testbank.db")
    conn.row_factory = sqlite3.Row
    return conn

def shuffle_and_relabel_choices(choices_list, correct_answer_letter):
    """Shuffle choices and reassign labels A, B, C, D, updating correct answer"""
    # Convert sqlite3.Row objects to dictionaries for modification
    choices = [dict(choice) for choice in choices_list]
    
    # Track which choice object is the correct one
    correct_choice_obj = None
    for choice in choices:
        if choice['choice_label'] == correct_answer_letter:
            correct_choice_obj = choice
            break
    
    # Shuffle the choices
    random.shuffle(choices)
    
    # Reassign labels A, B, C, D and find new correct answer
    new_correct_answer = None
    for idx, choice in enumerate(choices):
        new_label = chr(65 + idx)  # A=65, B=66, C=67, D=68
        choice['choice_label'] = new_label
        if choice is correct_choice_obj:
            new_correct_answer = new_label
    
    return choices, new_correct_answer

def sanitize_filename(filename):
    """Remove or replace invalid filename characters"""
    invalid_chars = '<>:"/\\|?*'
    for char in invalid_chars:
        filename = filename.replace(char, '_')
    return filename.strip()

def save_exam_to_history(conn, chapter_id, filename, file_path, num_questions, include_problems):
    """Record generated exam in exam_history table"""
    cursor = conn.cursor()
    # Use local system time instead of UTC
    created_date = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    cursor.execute("""
        INSERT INTO exam_history (chapter_id, filename, file_path, created_date, num_questions, include_problems)
        VALUES (?, ?, ?, ?, ?, ?)
    """, (chapter_id, filename, str(file_path), created_date, num_questions, 1 if include_problems else 0))
    conn.commit()
    return cursor.lastrowid

def generate_docx_exam(chapter, questions_with_choices, include_problems):
    """Generate exam as DOCX document"""
    doc = Document()
    
    # Title and chapter info
    title = doc.add_heading(f"Chapter {chapter['chapter_number']}: {chapter['chapter_title']}", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    exam_heading = doc.add_heading('Examination', level=2)
    exam_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Student info section
    student_info = doc.add_paragraph()
    student_info.add_run('Student Name: ').bold = True
    student_info.add_run('_' * 50)
    student_info.add_run('     Date: ').bold = True
    student_info.add_run('_' * 20)
    
    doc.add_paragraph()  # Blank line
    
    # Questions
    for idx, (question, choices) in enumerate(questions_with_choices, 1):
        # Problem text if included
        if include_problems and question.get('problem_text'):
            problem_para = doc.add_paragraph()
            problem_run = problem_para.add_run(f'Problem: {question["problem_text"]}')
            problem_run.italic = True
            problem_run.font.color.rgb = RGBColor(13, 115, 119)  # Teal color
        
        # Question text
        question_para = doc.add_paragraph(style='List Number')
        question_para.paragraph_format.left_indent = Inches(0)
        question_para.clear()
        question_run = question_para.add_run(f'{idx}. {question["question_text"]}')
        question_run.bold = True
        
        # Choice options
        for choice in choices:
            choice_para = doc.add_paragraph(
                f'{choice["choice_label"]}. {choice["choice_text"]}',
                style='List Bullet'
            )
            choice_para.paragraph_format.left_indent = Inches(0.5)
        
        doc.add_paragraph()  # Blank line between questions
    
    return doc

@app.route("/")
def index():
    conn = get_db()
    cursor = conn.cursor()
    
    # Get counts
    cursor.execute("SELECT COUNT(*) as count FROM chapters")
    chapters_count = cursor.fetchone()["count"]
    
    cursor.execute("SELECT COUNT(*) as count FROM questions")
    questions_count = cursor.fetchone()["count"]
    
    cursor.execute("SELECT COUNT(*) as count FROM problems")
    problems_count = cursor.fetchone()["count"]
    
    conn.close()
    
    return render_template("dashboard.html", 
                         chapters_count=chapters_count,
                         questions_count=questions_count,
                         problems_count=problems_count)

@app.route("/add_question")
def add_question():
    conn = get_db()
    cursor = conn.cursor()
    
    # Get all existing chapters
    cursor.execute("SELECT * FROM chapters ORDER BY chapter_number")
    chapters = cursor.fetchall()
    
    # Get all existing problems with their question counts and chapter info
    cursor.execute("""
        SELECT 
            p.id,
            p.problem_code,
            p.chapter_id,
            c.chapter_number,
            c.chapter_title,
            COUNT(q.id) as question_count
        FROM problems p
        JOIN chapters c ON p.chapter_id = c.id
        LEFT JOIN questions q ON p.id = q.problem_id
        GROUP BY p.id
        ORDER BY c.chapter_number, p.problem_code
    """)
    existing_problems = cursor.fetchall()
    
    conn.close()
    
    return render_template("add_question.html", chapters=chapters, existing_problems=existing_problems)

@app.route("/chapters")
def chapters():
    conn = get_db()
    cursor = conn.cursor()
    cursor.execute("SELECT * FROM chapters")
    chapters_list = cursor.fetchall()
    
    # Get problem count for each chapter
    problems_by_chapter = {}
    for chapter in chapters_list:
        cursor.execute("SELECT * FROM problems WHERE chapter_id = ?", (chapter["id"],))
        problems_by_chapter[chapter["id"]] = cursor.fetchall()
    
    conn.close()

    return render_template("chapters.html", 
                         chapters=chapters_list,
                         problems_by_chapter=problems_by_chapter)


@app.route("/delete_chapter/<int:chapter_id>", methods=["POST"])
def delete_chapter(chapter_id):
    """Delete a chapter and all associated problems, questions, and choices."""
    conn = get_db()
    cursor = conn.cursor()
    try:
        # Delete choices linked to questions under problems in this chapter
        cursor.execute("""
            DELETE FROM choices WHERE question_id IN (
                SELECT q.id FROM questions q
                JOIN problems p ON q.problem_id = p.id
                WHERE p.chapter_id = ?
            )
        """, (chapter_id,))

        # Delete questions under problems in this chapter
        cursor.execute("""
            DELETE FROM questions WHERE problem_id IN (
                SELECT id FROM problems WHERE chapter_id = ?
            )
        """, (chapter_id,))

        # Delete problems under this chapter
        cursor.execute("DELETE FROM problems WHERE chapter_id = ?", (chapter_id,))

        # Finally delete the chapter
        cursor.execute("DELETE FROM chapters WHERE id = ?", (chapter_id,))

        conn.commit()
    except Exception:
        conn.rollback()
    finally:
        conn.close()

    return redirect(url_for('chapters'))

@app.route("/quiz/<int:chapter_id>")
def quiz(chapter_id):
    conn = get_db()
    cursor = conn.cursor()

    # Get all problems under this chapter
    cursor.execute("""
        SELECT * FROM problems
        WHERE chapter_id = ?
    """, (chapter_id,))
    problems = cursor.fetchall()

    full_questions = []

    for problem in problems:
        # Get all questions under this problem
        cursor.execute("""
            SELECT * FROM questions
            WHERE problem_id = ?
        """, (problem["id"],))
        questions = cursor.fetchall()

        for q in questions:
            # Get choices for each question
            cursor.execute("""
                SELECT * FROM choices
                WHERE question_id = ?
            """, (q["id"],))
            choices = cursor.fetchall()

            full_questions.append({
                "problem": problem,
                "question": q,
                "choices": choices
            })

    conn.close()

    return render_template("quiz.html", questions=full_questions)

@app.route("/submit", methods=["POST"])
def submit():
    conn = get_db()
    cursor = conn.cursor()

    score = 0
    results = []

    for key in request.form:
        question_id = key.split("_")[1]
        selected = request.form[key]

        cursor.execute("SELECT * FROM questions WHERE id = ?", (question_id,))
        question = cursor.fetchone()

        is_correct = selected == question["correct_choice"]
        if is_correct:
            score += 1

        results.append({
            "question_text": question["question_text"],
            "selected": selected,
            "correct": question["correct_choice"],
            "explanation": question["explanation"],
            "is_correct": is_correct
        })

    conn.close()

    return render_template("result.html", score=score, results=results)

@app.route("/save", methods=["POST"])
def save_question():
    chapter_source = request.form.get("chapter_source", "existing")
    problem_source = request.form.get("problem_source", "new")
    
    # Get question data (same for both cases)
    question_text = request.form["question_text"]
    choice_a = request.form["choice_a"]
    choice_b = request.form["choice_b"]
    choice_c = request.form["choice_c"]
    choice_d = request.form["choice_d"]
    correct_choice = request.form["correct_choice"]
    explanation = request.form.get("explanation", "")

    conn = get_db()
    cursor = conn.cursor()

    # Step 1: Determine chapter_id
    if chapter_source == "existing":
        chapter_id = request.form["chapter_id"]
    else:
        # Creating new chapter - check if chapter number already exists
        new_chapter_number = request.form["new_chapter_number"]
        new_chapter_title = request.form["new_chapter_title"]
        
        cursor.execute("SELECT id FROM chapters WHERE chapter_number = ?", (new_chapter_number,))
        result = cursor.fetchone()
        
        if result:
            conn.close()
            return render_template("error.html", 
                                 error_message=f"Chapter number {new_chapter_number} already exists. Please use a different chapter number or select an existing chapter.")
        
        # Create new chapter
        cursor.execute(
            "INSERT INTO chapters (chapter_number, chapter_title) VALUES (?, ?)",
            (new_chapter_number, new_chapter_title)
        )
        chapter_id = cursor.lastrowid

    # Step 2: Handle problem (new or existing)
    if problem_source == "new":
        # Create new problem
        section = request.form.get("section")
        problem_code = request.form["problem_code"]
        standard = request.form.get("standard")
        problem_text = request.form["problem_text"]

        cursor.execute("""
            INSERT INTO problems (chapter_id, section, problem_code, standard, problem_text)
            VALUES (?, ?, ?, ?, ?)
        """, (chapter_id, section, problem_code, standard, problem_text))

        problem_id = cursor.lastrowid
    else:
        # Use existing problem
        problem_id = request.form["existing_problem_id"]

    # 3️⃣ Insert question
    cursor.execute("""
        INSERT INTO questions (problem_id, question_text, correct_choice, explanation)
        VALUES (?, ?, ?, ?)
    """, (problem_id, question_text, correct_choice, explanation))

    question_id = cursor.lastrowid

    # 4️⃣ Insert choices
    choices = [
        ("A", choice_a),
        ("B", choice_b),
        ("C", choice_c),
        ("D", choice_d),
    ]

    for label, text in choices:
        cursor.execute(
            "INSERT INTO choices (question_id, choice_label, choice_text) VALUES (?, ?, ?)",
            (question_id, label, text)
        )

    conn.commit()
    conn.close()

    return redirect("/")

@app.route("/chapter/<int:chapter_id>")
def chapter_detail(chapter_id):
    conn = get_db()
    cursor = conn.cursor()
    
    # Get search/filter parameters
    search_query = request.args.get('search', '').strip().lower()
    section_filter = request.args.get('section', '').strip()
    standard_filter = request.args.get('standard', '').strip()
    
    # Get chapter
    cursor.execute("SELECT * FROM chapters WHERE id = ?", (chapter_id,))
    chapter = cursor.fetchone()
    
    # Get problems and convert to dictionaries
    cursor.execute("SELECT * FROM problems WHERE chapter_id = ?", (chapter_id,))
    all_problems = [dict(row) for row in cursor.fetchall()]
    
    # Get unique sections and standards for filter dropdowns
    cursor.execute("SELECT DISTINCT section FROM problems WHERE chapter_id = ? AND section IS NOT NULL ORDER BY section", (chapter_id,))
    sections = [row[0] for row in cursor.fetchall()]
    
    cursor.execute("SELECT DISTINCT standard FROM problems WHERE chapter_id = ? AND standard IS NOT NULL ORDER BY standard", (chapter_id,))
    standards = [row[0] for row in cursor.fetchall()]
    
    # Get questions and choices for each problem
    questions_by_problem = {}
    choices_by_question = {}
    
    for problem in all_problems:
        cursor.execute("SELECT * FROM questions WHERE problem_id = ?", (problem["id"],))
        questions = [dict(row) for row in cursor.fetchall()]
        questions_by_problem[problem["id"]] = questions
        
        for question in questions:
            cursor.execute("SELECT * FROM choices WHERE question_id = ?", (question["id"],))
            choices = [dict(row) for row in cursor.fetchall()]
            choices_by_question[question["id"]] = choices
    
    # Apply filters
    filtered_problems = []
    for problem in all_problems:
        # Apply section filter
        if section_filter and problem.get('section') != section_filter:
            continue
        
        # Apply standard filter
        if standard_filter and problem.get('standard') != standard_filter:
            continue
        
        # Apply search filter (search in problem text and question text)
        if search_query:
            has_match = False
            if search_query in problem.get('problem_text', '').lower():
                has_match = True
            else:
                for question in questions_by_problem.get(problem["id"], []):
                    if search_query in question.get('question_text', '').lower():
                        has_match = True
                        break
            
            if not has_match:
                continue
        
        filtered_problems.append(problem)
    
    conn.close()
    
    return render_template("chapter_detail.html", 
                         chapter=chapter, 
                         problems=filtered_problems,
                         questions_by_problem=questions_by_problem,
                         choices_by_question=choices_by_question,
                         sections=sections,
                         standards=standards,
                         search_query=search_query,
                         section_filter=section_filter,
                         standard_filter=standard_filter)

@app.route("/api/chapter/<int:chapter_id>/questions")
def api_get_chapter_questions(chapter_id):
    """API endpoint to fetch all questions for a chapter (for manual selection)"""
    conn = get_db()
    cursor = conn.cursor()
    
    # Get all problems for the chapter
    cursor.execute("SELECT id, problem_code, section, standard, problem_text FROM problems WHERE chapter_id = ? ORDER BY problem_code", (chapter_id,))
    problems = cursor.fetchall()
    
    questions_data = []
    for problem in problems:
        cursor.execute("SELECT id, question_text FROM questions WHERE problem_id = ? ORDER BY id", (problem['id'],))
        questions = cursor.fetchall()
        
        for question in questions:
            questions_data.append({
                'id': question['id'],
                'question_text': question['question_text'],
                'problem_code': problem['problem_code'],
                'section': problem['section'] or '',
                'standard': problem['standard'] or ''
            })
    
    conn.close()
    return jsonify({'questions': questions_data})

@app.route("/edit_question/<int:question_id>")
def edit_question(question_id):
    conn = get_db()
    cursor = conn.cursor()
    
    # Get question
    cursor.execute("SELECT * FROM questions WHERE id = ?", (question_id,))
    question = cursor.fetchone()
    
    # Get problem
    cursor.execute("SELECT * FROM problems WHERE id = ?", (question["problem_id"],))
    problem = cursor.fetchone()
    
    # Get chapter
    cursor.execute("SELECT * FROM chapters WHERE id = ?", (problem["chapter_id"],))
    chapter = cursor.fetchone()
    
    # Get choices
    cursor.execute("SELECT * FROM choices WHERE question_id = ? ORDER BY choice_label", (question_id,))
    choices = cursor.fetchall()
    
    conn.close()
    
    return render_template("edit_question.html",
                         question=question,
                         problem=problem,
                         chapter=chapter,
                         choices=choices)

@app.route("/save_edit/<int:question_id>", methods=["POST"])
def save_edit(question_id):
    problem_code = request.form["problem_code"]
    section = request.form.get("section")
    standard = request.form["standard"]
    problem_text = request.form["problem_text"]
    question_text = request.form["question_text"]
    choice_a = request.form["choice_A"]
    choice_b = request.form["choice_B"]
    choice_c = request.form["choice_C"]
    choice_d = request.form["choice_D"]
    correct_choice = request.form["correct_choice"]
    explanation = request.form["explanation"]
    
    conn = get_db()
    cursor = conn.cursor()
    
    # Get current question to find problem_id
    cursor.execute("SELECT problem_id FROM questions WHERE id = ?", (question_id,))
    problem_id = cursor.fetchone()["problem_id"]
    
    # Update problem
    cursor.execute("""
        UPDATE problems 
        SET problem_code = ?, section = ?, standard = ?, problem_text = ?
        WHERE id = ?
    """, (problem_code, section, standard, problem_text, problem_id))
    
    # Update question
    cursor.execute("""
        UPDATE questions
        SET question_text = ?, correct_choice = ?, explanation = ?
        WHERE id = ?
    """, (question_text, correct_choice, explanation, question_id))
    
    # Update choices
    choices = [
        ("A", choice_a),
        ("B", choice_b),
        ("C", choice_c),
        ("D", choice_d),
    ]
    
    for label, text in choices:
        cursor.execute("""
            UPDATE choices
            SET choice_text = ?
            WHERE question_id = ? AND choice_label = ?
        """, (text, question_id, label))
    
    conn.commit()
    
    # Get chapter_id for redirect
    cursor.execute("SELECT chapter_id FROM problems WHERE id = ?", (problem_id,))
    chapter_id = cursor.fetchone()["chapter_id"]
    
    conn.close()
    
    return redirect(f"/chapter/{chapter_id}")

@app.route("/delete_question/<int:question_id>")
def delete_question(question_id):
    conn = get_db()
    cursor = conn.cursor()
    
    # Get problem_id and chapter_id
    cursor.execute("SELECT problem_id FROM questions WHERE id = ?", (question_id,))
    problem_id = cursor.fetchone()["problem_id"]
    
    cursor.execute("SELECT chapter_id FROM problems WHERE id = ?", (problem_id,))
    chapter_id = cursor.fetchone()["chapter_id"]
    
    # Delete choices
    cursor.execute("DELETE FROM choices WHERE question_id = ?", (question_id,))
    
    # Delete question
    cursor.execute("DELETE FROM questions WHERE id = ?", (question_id,))
    
    # Check if problem has any questions left, if not delete it
    cursor.execute("SELECT COUNT(*) as count FROM questions WHERE problem_id = ?", (problem_id,))
    count = cursor.fetchone()["count"]
    if count == 0:
        cursor.execute("DELETE FROM problems WHERE id = ?", (problem_id,))
    
    conn.commit()
    conn.close()
    
    return redirect(f"/chapter/{chapter_id}")

@app.route("/generate_exam")
def generate_exam():
    import json
    # Accept optional query params to prefill the form when returning from preview
    selected_chapter_id = request.args.get('chapter_id')
    selected_num_questions = request.args.get('num_questions')
    selected_include_problems = request.args.get('include_problems')
    selected_custom_filename = request.args.get('custom_filename')
    selected_file_format = request.args.get('file_format', 'pdf')

    conn = get_db()
    cursor = conn.cursor()
    
    # Get all chapters
    cursor.execute("SELECT * FROM chapters ORDER BY chapter_number")
    chapters = cursor.fetchall()
    
    # For each chapter, count questions
    question_counts = {}
    for chapter in chapters:
        cursor.execute("""
            SELECT COUNT(*) as count FROM questions q
            JOIN problems p ON q.problem_id = p.id
            WHERE p.chapter_id = ?
        """, (chapter["id"],))
        question_counts[chapter["id"]] = cursor.fetchone()["count"]
    
    conn.close()

    return render_template("generate_exam.html", 
                         chapters=chapters,
                         question_counts=question_counts,
                         question_counts_json=json.dumps(question_counts),
                         selected_chapter_id=selected_chapter_id,
                         selected_num_questions=selected_num_questions,
                         selected_include_problems=selected_include_problems,
                         selected_custom_filename=selected_custom_filename,
                         selected_file_format=selected_file_format)

@app.route("/preview_exam", methods=["POST"])
def preview_exam():
    """Preview exam or download as PDF or DOCX"""
    import json
    
    action = request.form.get("action", "preview")
    chapter_id = request.form.get("chapter_id")
    num_questions = int(request.form.get("num_questions", 10))
    include_problems = request.form.get("include_problems") == "yes"
    custom_filename = request.form.get("custom_filename", "").strip()
    file_format = request.form.get("file_format", "pdf").lower()
    mode = request.form.get("mode", "random")
    selected_question_ids = request.form.get("selected_question_ids", "")
    
    conn = get_db()
    cursor = conn.cursor()
    
    # Get chapter info
    cursor.execute("SELECT * FROM chapters WHERE id = ?", (chapter_id,))
    chapter = dict(cursor.fetchone())
    
    # Check if question_ids were passed (from preview download)
    question_ids_str = request.form.get("question_ids", "")
    
    if question_ids_str and action == "download":
        # Reuse the exact questions and choices from preview by fetching them in order
        question_ids = [int(qid) for qid in question_ids_str.split(',')]
        documents_to_use = []
        
        for qid in question_ids:
            # Get the question with its problem text
            cursor.execute("""
                SELECT q.*, p.problem_text FROM questions q
                JOIN problems p ON q.problem_id = p.id
                WHERE q.id = ?
            """, (qid,))
            question = dict(cursor.fetchone())
            
            # Get the choices for this question
            cursor.execute("SELECT * FROM choices WHERE question_id = ? ORDER BY choice_label", (qid,))
            choices = cursor.fetchall()
            
            # Shuffle choices and reassign labels (same shuffling algorithm)
            shuffled_choices, new_correct = shuffle_and_relabel_choices(choices, question['correct_choice'])
            question['correct_choice'] = new_correct
            
            documents_to_use.append((question, shuffled_choices))
        
        all_questions_count = len(question_ids)
    else:
        documents_to_use = None
        all_questions_count = 0
    
    # If we don't have pre-selected questions, generate new ones based on mode
    if documents_to_use is None:
        # Get all questions from this chapter with their problem info
        cursor.execute("""
            SELECT q.*, p.problem_text FROM questions q
            JOIN problems p ON q.problem_id = p.id
            WHERE p.chapter_id = ?
        """, (chapter_id,))
        all_questions = cursor.fetchall()
        
        # Select questions based on mode
        if mode == "manual" and selected_question_ids:
            # User manually selected specific questions
            selected_ids = [int(qid) for qid in selected_question_ids.split(',')]
            selected_questions = [dict(q) for q in all_questions if q['id'] in selected_ids]
            # Reorder to match user's selection order
            selected_questions.sort(key=lambda q: selected_ids.index(q['id']))
        else:
            # Random selection
            selected_questions = random.sample(all_questions, min(num_questions, len(all_questions)))
            selected_questions = [dict(q) for q in selected_questions]
        
        # Get choices for each selected question
        questions_with_choices = []
        for question in selected_questions:
            cursor.execute("SELECT * FROM choices WHERE question_id = ? ORDER BY choice_label", (question["id"],))
            choices = cursor.fetchall()
            # Shuffle choices and reassign labels
            shuffled_choices, new_correct = shuffle_and_relabel_choices(choices, question['correct_choice'])
            # Update the question's correct answer to reflect new position
            question['correct_choice'] = new_correct
            questions_with_choices.append((question, shuffled_choices))
        
        documents_to_use = questions_with_choices
        all_questions_count = len(all_questions)
    
    # If preview requested, render preview template
    if action == "preview":
        # Prepare question IDs to pass back (in order)
        question_ids = ','.join(str(q[0]['id']) for q in documents_to_use)
        
        conn.close()
        return render_template("preview_exam.html",
                             chapter=chapter,
                             questions_with_choices=documents_to_use,
                             num_questions=num_questions,
                             total_available=all_questions_count,
                             include_problems=include_problems,
                             custom_filename=custom_filename,
                             file_format=file_format,
                             question_ids=question_ids,
                             mode=mode,
                             selected_question_ids=selected_question_ids)
    
    # Otherwise, generate PDF or DOCX based on file_format
    # Determine filename first so we can use it for the document
    if custom_filename:
        base_filename = custom_filename
    else:
        base_filename = f"exam_chapter_{chapter['chapter_number']}_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
    
    base_filename = sanitize_filename(base_filename)
    
    # Generate DOCX if requested
    if file_format == 'docx':
        docx_filename = base_filename if base_filename.endswith('.docx') else base_filename + '.docx'
        docx_path = EXAMS_DIR / docx_filename
        
        # Generate DOCX document
        doc = generate_docx_exam(chapter, documents_to_use, include_problems)
        doc.save(str(docx_path))
        
        # Save to exam history (use actual count of questions used)
        save_exam_to_history(conn, int(chapter_id), docx_filename, docx_path, len(documents_to_use), include_problems)
        conn.close()
        
        return send_file(
            str(docx_path),
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            as_attachment=True,
            download_name=docx_filename
        )
    
    # Generate PDF (default)
    pdf_filename = base_filename if base_filename.endswith('.pdf') else base_filename + '.pdf'
    pdf_path = EXAMS_DIR / pdf_filename
    
    # Build PDF document with the actual file path
    pdf = SimpleDocTemplate(str(pdf_path), pagesize=letter)
    styles = getSampleStyleSheet()
    story = []
    
    # Title
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=18,
        textColor=colors.HexColor('#2c3e50'),
        spaceAfter=12,
        alignment=1  # Center
    )
    story.append(Paragraph(f"Chapter {chapter['chapter_number']}: {chapter['chapter_title']}", title_style))
    story.append(Paragraph(f"Examination", styles['Heading2']))
    story.append(Spacer(1, 0.3*inch))
    
    # Student info (blank lines for writing)
    story.append(Paragraph("Student Name: ___________________________________     Date: _________________", styles['Normal']))
    story.append(Spacer(1, 0.2*inch))
    
    # Add questions to PDF
    for idx, (question, choices) in enumerate(documents_to_use, 1):
        # Add problem text if requested
        if include_problems and question.get('problem_text'):
            problem_style = ParagraphStyle(
                'Problem',
                parent=styles['Normal'],
                fontSize=10,
                textColor=colors.HexColor('#8e44ad'),
                spaceAfter=8,
                leftIndent=0
            )
            story.append(Paragraph(f"<i><b>Problem:</b> {question['problem_text']}</i>", problem_style))
        
        story.append(Paragraph(f"<b>Question {idx}:</b> {question['question_text']}", styles['Normal']))
        
        # Add choices
        for choice in choices:
            story.append(Paragraph(f"&nbsp;&nbsp;&nbsp;&nbsp;<b>{choice['choice_label']}.</b> {choice['choice_text']}", styles['Normal']))
        
        story.append(Spacer(1, 0.15*inch))
    
    # Page break for answer key
    story.append(PageBreak())
    story.append(Paragraph("ANSWER KEY", title_style))
    story.append(Spacer(1, 0.3*inch))
    
    # Create paragraph style for explanations (smaller font, left-aligned)
    explanation_style = ParagraphStyle(
        'Explanation',
        parent=styles['Normal'],
        fontSize=8,
        leading=10
    )
    
    # Add answer key with wrapped explanations
    answer_data = [["Question", "Correct Answer", "Explanation"]]
    for idx, (question, choices) in enumerate(documents_to_use, 1):
        answer_data.append([
            f"Q {idx}",
            question['correct_choice'],
            Paragraph(question['explanation'], explanation_style)
        ])
    
    # Create answer table
    answer_table = Table(answer_data, colWidths=[0.8*inch, 1.2*inch, 4.2*inch])
    answer_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#3498db')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (0, -1), 'CENTER'),
        ('ALIGN', (1, 0), (1, -1), 'CENTER'),
        ('ALIGN', (2, 0), (2, -1), 'LEFT'),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 11),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('TOPPADDING', (0, 1), (-1, -1), 8),
        ('BOTTOMPADDING', (0, 1), (-1, -1), 8),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('FONTSIZE', (0, 1), (-1, -1), 9),
    ]))
    story.append(answer_table)
    
    # Build the PDF file
    pdf.build(story)
    
    # Save to exam history (use actual count of questions used)
    save_exam_to_history(conn, int(chapter_id), pdf_filename, pdf_path, len(documents_to_use), include_problems)
    conn.close()
    
    return send_file(
        str(pdf_path),
        mimetype="application/pdf",
        as_attachment=True,
        download_name=pdf_filename
    )

@app.route("/export_exam_pdf", methods=["POST"])
def export_exam_pdf():
    """Export exam questions and answer key to PDF"""
    
    chapter_id = request.form.get("chapter_id")
    num_questions = int(request.form.get("num_questions", 10))
    
    conn = get_db()
    cursor = conn.cursor()
    
    # Get chapter info
    cursor.execute("SELECT * FROM chapters WHERE id = ?", (chapter_id,))
    chapter = dict(cursor.fetchone())
    
    # Get all questions from this chapter
    cursor.execute("""
        SELECT q.* FROM questions q
        JOIN problems p ON q.problem_id = p.id
        WHERE p.chapter_id = ?
    """, (chapter_id,))
    all_questions = cursor.fetchall()
    
    # Randomly select num_questions
    selected_questions = random.sample(all_questions, min(num_questions, len(all_questions)))
    
    # Convert questions to dictionaries for modification
    selected_questions = [dict(q) for q in selected_questions]
    
    # Get choices for each selected question
    questions_with_choices = []
    for question in selected_questions:
        cursor.execute("SELECT * FROM choices WHERE question_id = ? ORDER BY choice_label", (question["id"],))
        choices = cursor.fetchall()
        # Shuffle choices and reassign labels
        shuffled_choices, new_correct = shuffle_and_relabel_choices(choices, question['correct_choice'])
        # Update the question's correct answer to reflect new position
        question['correct_choice'] = new_correct
        questions_with_choices.append((question, shuffled_choices))
    
    # Build PDF document
    buffer = BytesIO()
    pdf = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []
    
    # Title
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=18,
        textColor=colors.HexColor('#2c3e50'),
        spaceAfter=12,
        alignment=1  # Center
    )
    story.append(Paragraph(f"Chapter {chapter['chapter_number']}: {chapter['chapter_title']}", title_style))
    story.append(Paragraph(f"Examination", styles['Heading2']))
    story.append(Spacer(1, 0.3*inch))
    
    # Student info (blank lines for writing)
    story.append(Paragraph("Student Name: ___________________________________     Date: _________________", styles['Normal']))
    story.append(Spacer(1, 0.2*inch))
    
    # Add questions to PDF
    for idx, (question, choices) in enumerate(questions_with_choices, 1):
        story.append(Paragraph(f"<b>Question {idx}:</b> {question['question_text']}", styles['Normal']))
        
        # Add choices
        for choice in choices:
            story.append(Paragraph(f"&nbsp;&nbsp;&nbsp;&nbsp;<b>{choice['choice_label']}.</b> {choice['choice_text']}", styles['Normal']))
        
        story.append(Spacer(1, 0.15*inch))
    
    # Page break for answer key
    story.append(PageBreak())
    story.append(Paragraph("ANSWER KEY", title_style))
    story.append(Spacer(1, 0.3*inch))
    
    # Create paragraph style for explanations (smaller font, left-aligned)
    explanation_style = ParagraphStyle(
        'Explanation',
        parent=styles['Normal'],
        fontSize=8,
        leading=10
    )

    # Add answer key with wrapped explanations
    answer_data = [["Question", "Correct Answer", "Explanation"]]
    for idx, (question, choices) in enumerate(questions_with_choices, 1):
        answer_data.append([
            f"Q {idx}",
            question['correct_choice'],
            Paragraph(question['explanation'], explanation_style)
        ])

    # Create answer table
    answer_table = Table(answer_data, colWidths=[0.8*inch, 1.2*inch, 4.2*inch])
    answer_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#3498db')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (0, -1), 'CENTER'),
        ('ALIGN', (1, 0), (1, -1), 'CENTER'),
        ('ALIGN', (2, 0), (2, -1), 'LEFT'),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 11),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('TOPPADDING', (0, 1), (-1, -1), 8),
        ('BOTTOMPADDING', (0, 1), (-1, -1), 8),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('FONTSIZE', (0, 1), (-1, -1), 9),
    ]))
    story.append(answer_table)
    
    conn.close()
    
    # Build PDF
    pdf.build(story)
    buffer.seek(0)
    
    return send_file(
        buffer,
        mimetype="application/pdf",
        as_attachment=True,
        download_name=f"exam_chapter_{chapter['chapter_number']}_{datetime.now().strftime('%Y%m%d')}.pdf"
    )

@app.route("/exam_history")
def exam_history():
    """View history of generated exams"""
    conn = get_db()
    cursor = conn.cursor()
    
    # Get all exam history with chapter info
    cursor.execute("""
        SELECT eh.*, c.chapter_number, c.chapter_title
        FROM exam_history eh
        JOIN chapters c ON eh.chapter_id = c.id
        ORDER BY eh.created_date DESC
    """)
    exams = cursor.fetchall()
    conn.close()
    
    return render_template("exam_history.html", exams=exams)

@app.route("/download_exam/<int:exam_id>")
def download_exam(exam_id):
    """Download a saved exam file"""
    conn = get_db()
    cursor = conn.cursor()
    
    # Get exam details
    cursor.execute("SELECT * FROM exam_history WHERE id = ?", (exam_id,))
    exam = cursor.fetchone()
    conn.close()
    
    if not exam:
        return "Exam not found", 404
    
    file_path = Path(exam['file_path'])
    if not file_path.exists():
        return "File not found", 404
    
    return send_file(
        str(file_path),
        mimetype="application/pdf",
        as_attachment=True,
        download_name=exam['filename']
    )

@app.route("/delete_exam/<int:exam_id>", methods=["POST"])
def delete_exam(exam_id):
    """Delete an exam from history and filesystem"""
    conn = get_db()
    cursor = conn.cursor()
    
    # Get exam details
    cursor.execute("SELECT * FROM exam_history WHERE id = ?", (exam_id,))
    exam = cursor.fetchone()
    
    if not exam:
        conn.close()
        return "Exam not found", 404
    
    # Delete file from disk
    file_path = Path(exam['file_path'])
    if file_path.exists():
        file_path.unlink()
    
    # Delete from database
    cursor.execute("DELETE FROM exam_history WHERE id = ?", (exam_id,))
    conn.commit()
    conn.close()
    
    return redirect("/exam_history")

if __name__ == "__main__":
    app.run(debug=True)